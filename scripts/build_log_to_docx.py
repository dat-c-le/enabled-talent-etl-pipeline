"""
Convert PIPELINE_BUILD_LOG.md into a formatted Word document.
Output: PIPELINE_BUILD_LOG.docx
"""
import re
from pathlib import Path

import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = Path("PIPELINE_BUILD_LOG.md")
DEST = Path("PIPELINE_BUILD_LOG.docx")

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x59, 0x59, 0x59)

_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline_runs(paragraph, text: str) -> None:
    """Parse **bold** and `code` spans and add styled runs to a paragraph."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        else:
            paragraph.add_run(part)


def add_shaded_cell(cell, text: str, bold=False, bg=None, font_color=None, size=10, align_center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline_runs(p, text)
    for run in p.runs:
        run.font.size = Pt(size)
        if bold:
            run.bold = True
        if font_color:
            run.font.color.rgb = font_color
    if bg:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), bg)
        cell._tc.get_or_add_tcPr().append(shd)


def parse_md_table(lines):
    """Parse a contiguous block of markdown table lines into rows of cell text."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(re.match(r"^-+$", c) for c in cells):
            continue  # separator row
        rows.append(cells)
    return rows


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = docx.Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    title_done = False

    while i < n:
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = NAVY
            p.space_after = Pt(6)
            title_done = True
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_heading(level=1)
            run = p.add_run(line[3:].strip())
            run.font.color.rgb = NAVY
            run.font.size = Pt(14)
            i += 1
            continue

        if line.strip() == "---":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "BFBFBF")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        if line.startswith("**Commit:**"):
            p = doc.add_paragraph()
            add_inline_runs(p, line)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = GRAY
            p.space_after = Pt(8)
            i += 1
            continue

        if line.startswith("| "):
            # Collect the whole table block
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_md_table(table_lines)
            if rows:
                table = doc.add_table(rows=0, cols=len(rows[0]))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"
                for r_idx, row_vals in enumerate(rows):
                    row = table.add_row()
                    for c_idx, val in enumerate(row_vals):
                        is_header = r_idx == 0
                        add_shaded_cell(
                            row.cells[c_idx], val,
                            bold=is_header,
                            bg="1F4E79" if is_header else ("EBF5FB" if r_idx % 2 == 0 else None),
                            font_color=RGBColor(0xFF, 0xFF, 0xFF) if is_header else None,
                        )
                doc.add_paragraph()
            continue

        if line.lstrip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line.lstrip()[2:])
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1

    doc.save(DEST)
    print(f"Saved: {DEST}")


if __name__ == "__main__":
    main()
